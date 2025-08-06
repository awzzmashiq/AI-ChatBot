import os
import re
import json
import threading
from flask import Flask, request, session, jsonify, make_response, Response, stream_template, redirect
from flask_cors import CORS
from dotenv import load_dotenv
import requests
import pytesseract
from PIL import Image
from faster_whisper import WhisperModel
import docx
import pandas as pd
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from storage_manager import storage_manager, GoogleDriveStorageProvider
from functools import wraps
from flask import request, Response
import os

# JWT and auth
from jose import JWTError, jwt
import datetime

# Load environment variables
load_dotenv()
THETA_API_KEY = os.getenv("THETA_API_KEY")
THETA_API_URL = os.getenv("THETA_API_URL", "https://ondemand.thetaedgecloud.com/infer_request/llama_3_1_70b/completions")
GOOGLE_CLIENT_ID = os.getenv("GOOGLE_CLIENT_ID")
# Secret key for JWT signing
JWT_SECRET = os.getenv("SECRET_KEY", "studybuddy_secret_key")

# Demo configuration
DEMO_MODE = os.getenv("DEMO_MODE", "false").lower() == "true"
DEMO_CODES = [
    os.getenv("DEMO_CODE_1", "INVESTOR_ACCESS_2024"),
    os.getenv("DEMO_CODE_2", "POC_SHOWCASE_KEY"), 
    os.getenv("DEMO_CODE_3", "STARTUP_DEMO_PASS")
]

# Basic authentication credentials
BASIC_AUTH_USERNAME = os.getenv('BASIC_AUTH_USERNAME', 'admin')
BASIC_AUTH_PASSWORD = os.getenv('BASIC_AUTH_PASSWORD', 'your-secure-password')

def check_auth(username, password):
    """Check if username/password combination is valid"""
    return username == BASIC_AUTH_USERNAME and password == BASIC_AUTH_PASSWORD

def authenticate():
    """Send a 401 response that enables basic auth"""
    return Response(
        'Could not verify your access level for that URL.\n'
        'You have to login with proper credentials', 401,
        {'WWW-Authenticate': 'Basic realm="Login Required"'})

def requires_auth(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        auth = request.authorization
        if not auth or not check_auth(auth.username, auth.password):
            return authenticate()
        return f(*args, **kwargs)
    return decorated

# Apply basic auth to all routes except login/signup
@app.before_request
def require_basic_auth():
    # Skip auth for login/signup endpoints
    if request.endpoint in ['login', 'signup', 'google_login']:
        return
    
    # Skip auth for static files and OPTIONS requests
    if request.method == 'OPTIONS' or request.path.startswith('/static/'):
        return
    
    auth = request.authorization
    if not auth or not check_auth(auth.username, auth.password):
        return authenticate()

# Check if API key is available
if not THETA_API_KEY:
    print("[WARNING] THETA_API_KEY not found in environment variables!")
    print("[WARNING] LLM functionality will not work without a valid API key.")
    print("[INFO] Please create a .env file with your THETA_API_KEY")

# Flask app setup
app = Flask(__name__)
app.secret_key = JWT_SECRET  # Secret key for session (also used for JWT signing)
# Enable CORS for cross-origin requests with credentials (allow React dev origin)
CORS(app, supports_credentials=True, resources={r"/*": {"origins": ["http://localhost:3000", "http://127.0.0.1:3000"]}})
print("[Flask] CORS configured for localhost:3000")

# Global data structures and paths
users_db_path = "users.json"
vectorstores_dir = "vectorstores"
vectorstores_cache = {}  # cache to store loaded vectorstores per user in memory
global_vectorstore = None  # global vectorstore for shared knowledge

books_dir = "books"
conversation_histories = {}  # in-memory chat history per user
chat_sessions = {}  # in-memory chat sessions per user

# Initialize embeddings model
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/multi-qa-mpnet-base-dot-v1")

# Ensure base directories exist
os.makedirs(vectorstores_dir, exist_ok=True)
os.makedirs(books_dir, exist_ok=True)

# Load or initialize user database
def load_users():
    if os.path.exists(users_db_path):
        with open(users_db_path, "r") as f:
            return json.load(f)
    return {}

def save_users(users):
    with open(users_db_path, "w") as f:
        json.dump(users, f)

users = load_users()

# Utility: sanitize username for file paths (replace non-alphanumeric chars)
def safe_filename(name):
    return re.sub(r'[^A-Za-z0-9_-]+', '_', name)

# Load or create a user's vectorstore from disk
def load_vectorstore_for_user(username):
    # Return cached vectorstore if already loaded
    if username in vectorstores_cache:
        return vectorstores_cache[username]
    user_vector_dir = os.path.join(vectorstores_dir, safe_filename(username))
    if os.path.isdir(user_vector_dir):
        try:
            vs = FAISS.load_local(user_vector_dir, embeddings)
            print(f"[INFO] Loaded existing vectorstore for {username}")
            vectorstores_cache[username] = vs
            return vs
        except Exception as e:
            print(f"[WARNING] Could not load vectorstore for {username}: {e}")
    return None

def save_vectorstore_for_user(username, vs):
    user_vector_dir = os.path.join(vectorstores_dir, safe_filename(username))
    os.makedirs(user_vector_dir, exist_ok=True)
    vs.save_local(user_vector_dir)
    # Cache the vectorstore in memory as well
    vectorstores_cache[username] = vs

# Global vectorstore functions for shared knowledge
def load_global_vectorstore():
    """Load or create global vectorstore for shared knowledge"""
    global global_vectorstore
    global_vector_dir = os.path.join(vectorstores_dir, "global")
    
    if os.path.isdir(global_vector_dir):
        try:
            global_vectorstore = FAISS.load_local(global_vector_dir, embeddings)
            print(f"[INFO] Loaded existing global vectorstore")
            return global_vectorstore
        except Exception as e:
            print(f"[WARNING] Could not load global vectorstore: {e}")
    
    # Create new global vectorstore if it doesn't exist
    print("[INFO] Creating new global vectorstore")
    global_vectorstore = FAISS.from_texts(["Welcome to Study Buddy! I'm here to help you with your studies."], embeddings)
    save_global_vectorstore()
    return global_vectorstore

def save_global_vectorstore():
    """Save global vectorstore to disk"""
    global global_vectorstore
    if global_vectorstore:
        global_vector_dir = os.path.join(vectorstores_dir, "global")
        os.makedirs(global_vector_dir, exist_ok=True)
        global_vectorstore.save_local(global_vector_dir)
        print("[INFO] Saved global vectorstore")

def update_global_vectorstore(texts, metadata=None):
    """Add new texts to global vectorstore"""
    global global_vectorstore
    if not global_vectorstore:
        global_vectorstore = load_global_vectorstore()
    
    # Add new texts to the vectorstore
    if metadata is None:
        metadata = [{"source": "global_knowledge", "timestamp": datetime.datetime.now().isoformat()} for _ in texts]
    
    global_vectorstore.add_texts(texts, metadata)
    save_global_vectorstore()
    print(f"[INFO] Updated global vectorstore with {len(texts)} new texts")

def search_global_knowledge(query, k=5):
    """Search global knowledge base"""
    global global_vectorstore
    if not global_vectorstore:
        global_vectorstore = load_global_vectorstore()
    
    try:
        results = global_vectorstore.similarity_search(query, k=k)
        return results
    except Exception as e:
        print(f"[ERROR] Global knowledge search failed: {e}")
        return []

# Document management functions
def get_user_documents(user):
    """Get list of all documents uploaded by user"""
    try:
        storage_provider = storage_manager.get_user_storage_provider(user)
        documents = storage_provider.list_files(user)
        
        # For Google Drive (Demo), return all documents since they're already "indexed"
        if isinstance(storage_provider, GoogleDriveStorageProvider):
            return documents
        
        # For local storage, filter to only show indexed files (maintain compatibility)
        user_dir = os.path.join(books_dir, safe_filename(user))
        indexed_list_path = os.path.join(user_dir, "indexed_files.json")
        
        indexed_files = []
        if os.path.exists(indexed_list_path):
            try:
                with open(indexed_list_path, "r") as idxf:
                    data = json.load(idxf)
                    if isinstance(data, list):
                        indexed_files = data
                    elif isinstance(data, dict) and "indexed_files" in data:
                        indexed_files = data["indexed_files"]
            except:
                indexed_files = []
        
        # Only return documents that are indexed
        indexed_documents = [doc for doc in documents if doc["filename"] in indexed_files]
        return indexed_documents
        
    except Exception as e:
        print(f"[Get User Documents Error] {e}")
        return []

def delete_user_document(user, filename):
    """Delete a document and remove it from vectorstore"""
    try:
        storage_provider = storage_manager.get_user_storage_provider(user)
        
        # Check if file exists in storage
        if not storage_provider.file_exists(user, filename):
            return False, "File not found"
        
        # Remove file from storage
        if not storage_provider.delete_file(user, filename):
            return False, "Failed to delete file from storage"
        
        # For Google Drive (Demo), no need to update indexed files list
        if not isinstance(storage_provider, GoogleDriveStorageProvider):
            # Update indexed files list (maintain compatibility for local storage)
            user_dir = os.path.join(books_dir, safe_filename(user))
            indexed_list_path = os.path.join(user_dir, "indexed_files.json")
            
            indexed_files = []
            if os.path.exists(indexed_list_path):
                try:
                    with open(indexed_list_path, "r") as idxf:
                        data = json.load(idxf)
                        if isinstance(data, list):
                            indexed_files = data
                        elif isinstance(data, dict) and "indexed_files" in data:
                            indexed_files = data["indexed_files"]
                except:
                    indexed_files = []
            
            # Remove filename from indexed list
            if filename in indexed_files:
                indexed_files.remove(filename)
                
                # Save updated list
                try:
                    with open(indexed_list_path, 'w') as idxf:
                        json.dump({"indexed_files": indexed_files}, idxf)
                except Exception as e:
                    print(f"[Indexed files save error] {e}")
        
        # Rebuild vectorstore without this document
        rebuild_user_vectorstore(user)
        
        return True, "Document deleted successfully"
        
    except Exception as e:
        print(f"[Document deletion error] {e}")
        return False, f"Error deleting document: {str(e)}"

def rebuild_user_vectorstore(user):
    """Rebuild user's vectorstore from remaining documents"""
    user_dir = os.path.join(books_dir, safe_filename(user))
    indexed_list_path = os.path.join(user_dir, "indexed_files.json")
    
    # Load indexed files list
    indexed_files = []
    if os.path.exists(indexed_list_path):
        try:
            with open(indexed_list_path, "r") as idxf:
                data = json.load(idxf)
                if isinstance(data, list):
                    indexed_files = data
                elif isinstance(data, dict) and "indexed_files" in data:
                    indexed_files = data["indexed_files"]
        except:
            indexed_files = []
    
    if not indexed_files:
        # No documents left, clear vectorstore
        if user in vectorstores_cache:
            del vectorstores_cache[user]
        user_vector_dir = os.path.join(vectorstores_dir, safe_filename(user))
        if os.path.exists(user_vector_dir):
            import shutil
            shutil.rmtree(user_vector_dir)
        return
    
    # Rebuild vectorstore from remaining documents
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    all_docs = []
    
    for filename in indexed_files:
        file_path = os.path.join(user_dir, filename)
        if not os.path.exists(file_path):
            continue
            
        try:
            ext = filename.lower().rsplit('.', 1)[-1]
            if ext == "pdf":
                loader = PyPDFLoader(file_path)
                pages = loader.load()
                all_docs.extend(splitter.split_documents(pages))
            elif ext in ["jpg", "jpeg", "png"]:
                text = pytesseract.image_to_string(Image.open(file_path))
                all_docs.extend(splitter.create_documents([text]))
            elif ext in ["mp3", "wav", "m4a"]:
                model = WhisperModel("base")
                segments, info = model.transcribe(file_path)
                result_text = " ".join([seg.text for seg in segments])
                all_docs.extend(splitter.create_documents([result_text]))
            elif ext == "docx":
                doc_obj = docx.Document(file_path)
                full_text = "\n".join([para.text for para in doc_obj.paragraphs])
                all_docs.extend(splitter.create_documents([full_text]))
            elif ext in ["xlsx", "xls"]:
                df = pd.read_excel(file_path, engine="openpyxl" if ext == "xlsx" else "xlrd")
                csv_text = df.to_csv(index=False)
                all_docs.extend(splitter.create_documents([csv_text]))
        except Exception as e:
            print(f"[Document processing error for {filename}] {e}")
            continue
    
    # Create new vectorstore
    if all_docs:
        try:
            vs = FAISS.from_documents(all_docs, embeddings)
            save_vectorstore_for_user(user, vs)
            print(f"[INFO] Rebuilt vectorstore for {user} with {len(all_docs)} documents")
        except Exception as e:
            print(f"[Vectorstore rebuild error] {e}")
    else:
        # No valid documents, clear vectorstore
        if user in vectorstores_cache:
            del vectorstores_cache[user]
        user_vector_dir = os.path.join(vectorstores_dir, safe_filename(user))
        if os.path.exists(user_vector_dir):
            import shutil
            shutil.rmtree(user_vector_dir)

def search_user_documents(user, query, k=5):
    """Search user's documents with query"""
    vs = load_vectorstore_for_user(user)
    if not vs:
        return []
    
    try:
        results = vs.similarity_search(query, k=k)
        return results
    except Exception as e:
        print(f"[Document search error] {e}")
        return []

# Initialize embedding model for document vectors (multilingual support)
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/multi-qa-mpnet-base-dot-v1")

# Theta API LLM interface (LLaMA 3 70B)
class ThetaLLM:
    def __init__(self, api_key, url=THETA_API_URL, temperature=0.5, top_p=0.7, max_tokens=500):
        self.url = url
        self.headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        self.temperature = temperature
        self.top_p = top_p
        self.max_tokens = max_tokens

    def invoke(self, messages):
        # messages: list of {"role": ..., "content": ...}
        payload = {
            "input": {
                "messages": messages,
                "temperature": self.temperature,
                "top_p": self.top_p,
                "max_tokens": self.max_tokens,
                "stream": False
            }
        }
        try:
            response = requests.post(self.url, headers=self.headers, json=payload, timeout=30)
            response.raise_for_status()
        except requests.exceptions.Timeout:
            print("[Theta API Timeout] Request timed out")
            return "⚠️ Request timed out. Please try again."
        except Exception as e:
            print(f"[Theta API Error] {e}")
            return "⚠️ Error from Theta API"
        data = response.json()
        try:
            assistant_msg = data["body"]["infer_requests"][0]["output"]["message"]
            if isinstance(assistant_msg, dict):
                return assistant_msg.get("content", "")
            return assistant_msg  # if it's already a string
        except Exception as e:
            print(f"[Theta API Error] {e}")
            return "⚠️ Error from Theta API"



# Initialize the LLM with API key
llm = ThetaLLM(api_key=THETA_API_KEY)

# Chat sessions management
def get_user_sessions(user):
    """Get all chat sessions for a user"""
    if user not in chat_sessions:
        filename = f"chat_sessions_{safe_filename(user)}.json"
        if os.path.exists(filename):
            try:
                with open(filename, "r") as f:
                    chat_sessions[user] = json.load(f)
            except:
                chat_sessions[user] = []
        else:
            # Create default session for existing users
            chat_sessions[user] = [{
                "id": "default",
                "name": "New Chat",
                "created_at": datetime.datetime.now().isoformat(),
                "updated_at": datetime.datetime.now().isoformat()
            }]
            # Save the default session to disk
            save_user_sessions(user)
    return chat_sessions[user]

def save_user_sessions(user):
    """Save chat sessions for a user"""
    sessions = chat_sessions.get(user, [])
    filename = f"chat_sessions_{safe_filename(user)}.json"
    with open(filename, "w") as f:
        json.dump(sessions, f, indent=2)

def get_conversation(user, session_id="default"):
    """Get conversation history for a specific chat session"""
    key = f"{user}_{session_id}"
    if key not in conversation_histories:
        filename = f"chat_history_{safe_filename(user)}_{session_id}.json"
        if os.path.exists(filename):
            try:
                with open(filename, "r") as f:
                    conversation_histories[key] = json.load(f)
            except:
                conversation_histories[key] = []
        else:
            conversation_histories[key] = []
    return conversation_histories[key]

def save_conversation(user, session_id="default"):
    """Save conversation history for a specific chat session"""
    key = f"{user}_{session_id}"
    hist = conversation_histories.get(key, [])
    filename = f"chat_history_{safe_filename(user)}_{session_id}.json"
    with open(filename, "w") as f:
        json.dump(hist, f, indent=2)

def create_new_session(user, name="New Chat"):
    """Create a new chat session"""
    import uuid
    session_id = str(uuid.uuid4())
    sessions = get_user_sessions(user)
    
    new_session = {
        "id": session_id,
        "name": name,
        "created_at": datetime.datetime.now().isoformat(),
        "updated_at": datetime.datetime.now().isoformat()
    }
    
    sessions.append(new_session)
    chat_sessions[user] = sessions
    save_user_sessions(user)
    
    # Initialize empty conversation for new session
    conversation_histories[f"{user}_{session_id}"] = []
    
    return new_session

def update_session_name(user, session_id, new_name):
    """Update the name of a chat session"""
    sessions = get_user_sessions(user)
    for session in sessions:
        if session["id"] == session_id:
            session["name"] = new_name
            session["updated_at"] = datetime.datetime.now().isoformat()
            break
    
    chat_sessions[user] = sessions
    save_user_sessions(user)
    return True

def delete_session(user, session_id):
    """Delete a chat session and its history"""
    sessions = get_user_sessions(user)
    sessions = [s for s in sessions if s["id"] != session_id]
    chat_sessions[user] = sessions
    save_user_sessions(user)
    
    # Delete conversation history file
    key = f"{user}_{session_id}"
    if key in conversation_histories:
        del conversation_histories[key]
    
    filename = f"chat_history_{safe_filename(user)}_{session_id}.json"
    if os.path.exists(filename):
        os.remove(filename)
    
    return True

# Extract simple profile info (name, interests) from past user messages
def extract_user_profile(history, default_name="User"):
    profile = {"name": default_name, "interests": set()}
    for msg in history:
        if msg.get("role") == "user":
            text = msg["content"].lower()
            if "my name is" in text:
                part = text.split("my name is", 1)[1].strip()
                if part:
                    name = part.split()[0]
                    profile["name"] = name.capitalize()
            if "i like" in text or "i am interested in" in text:
                if "i like" in text:
                    interest = text.split("i like", 1)[1].strip()
                else:
                    interest = text.split("interested in", 1)[1].strip()
                if interest:
                    interest = interest.split('.')[0]
                    profile["interests"].add(interest.capitalize())
    return profile

# Helper function to generate JWT token
def create_jwt_for_user(username):
    # Set expiration for token (e.g., 24 hours from now)
    exp = datetime.datetime.utcnow() + datetime.timedelta(hours=24)
    payload = {"username": username, "exp": exp}
    token = jwt.encode(payload, JWT_SECRET, algorithm='HS256')
    return token

# Helper function to get current user from token in request
def get_user_from_token():
    token = request.cookies.get('token')
    if not token:
        print("[DEBUG] No token found in cookies")
        return None
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=['HS256'])
        username = payload.get('username')
        if not username:
            print("[DEBUG] No username found in token payload")
            return None
        print(f"[DEBUG] Valid token for user: {username}")
        return username
    except JWTError as e:
        # Token is invalid or expired
        print(f"[DEBUG] JWT decode error: {e}")
        return None
    except Exception as e:
        print(f"[DEBUG] Unexpected error in token validation: {e}")
        return None

# User authentication routes

@app.route("/api/signup", methods=["POST"])
def signup():
    data = request.get_json()
    if not data:
        return jsonify({"success": False, "message": "No data provided"}), 400
    username = data.get("username", "").strip()
    password = data.get("password", "").strip()
    if not username or not password:
        return jsonify({"success": False, "message": "Username and password required"}), 400
    if username in users:
        return jsonify({"success": False, "message": "User already exists"}), 409
    # Save new user (store plain password for demo purposes)
    users[username] = password
    save_users(users)
    # Initialize chat sessions for new user
    sessions = get_user_sessions(username)
    if not sessions:
        create_new_session(username, "New Chat")
    # Create JWT token for user and set it in an HTTP-only cookie
    token = create_jwt_for_user(username)
    resp = jsonify({"success": True, "message": "Account created", "user": username})
    secure_flag = False if app.debug else True  # use Secure cookies in production
    resp.set_cookie("token", token, httponly=True, secure=False, samesite='Lax')
    return resp

@app.route("/api/login", methods=["POST"])
def login():
    data = request.get_json()
    if not data:
        return jsonify({"success": False, "message": "No data provided"}), 400
    username = data.get("username", "").strip()
    password = data.get("password", "").strip()
    print(f"[DEBUG] Login attempt for username: {username}")
    if not username or not password:
        return jsonify({"success": False, "message": "Username and password required"}), 400
    if username not in users or users.get(username) != password:
        print(f"[DEBUG] Invalid credentials for {username}. Users in DB: {list(users.keys())}")
        return jsonify({"success": False, "message": "Invalid credentials"}), 401
    # User authenticated, prepare their sessions
    sessions = get_user_sessions(username)  # load existing sessions if any
    # Ensure user has at least one session
    if not sessions:
        create_new_session(username, "New Chat")
    token = create_jwt_for_user(username)
    print(f"[DEBUG] Created JWT token for {username}")
    resp = jsonify({"success": True, "message": "Login successful", "user": username})
    secure_flag = False if app.debug else True
    resp.set_cookie("token", token, httponly=True, secure=secure_flag, samesite='Lax')
    print(f"[DEBUG] Set cookie with token for {username}")
    return resp

@app.route("/api/google-login", methods=["POST"])
def google_login():
    data = request.get_json()
    token = data.get("credential")
    if not token or not GOOGLE_CLIENT_ID:
        return jsonify({"success": False, "message": "Missing Google token or server not configured"}), 400
    try:
        # Verify Google ID token
        from google.oauth2 import id_token as google_id_token
        from google.auth.transport import requests as google_requests
        idinfo = google_id_token.verify_oauth2_token(token, google_requests.Request(), GOOGLE_CLIENT_ID)
        email = idinfo.get("email")
        if not email:
            return jsonify({"success": False, "message": "Google login failed"}), 400
        username = email
        if username not in users:
            # New Google user, create an account with no password
            users[username] = ""
            save_users(users)
            sessions = get_user_sessions(username)  # Initialize sessions
            if not sessions:
                create_new_session(username, "New Chat")
        else:
            # Existing user, load their sessions
            sessions = get_user_sessions(username)
            if not sessions:
                create_new_session(username, "New Chat")
        # For Google users, automatically set Google Drive as storage preference
        try:
            storage_manager.set_user_storage_preference(username, 'google_drive')
            print(f"[Google Login] Auto-set Google Drive storage for user: {username}")
        except Exception as e:
            print(f"[Google Login] Failed to set Google Drive storage: {e}")
        
        # Issue JWT for the authenticated Google user
        jwt_token = create_jwt_for_user(username)
        resp = jsonify({"success": True, "message": "Google login successful", "user": username})
        secure_flag = False if app.debug else True
        resp.set_cookie("token", jwt_token, httponly=True, secure=secure_flag, samesite='Lax')
        return resp
    except Exception as e:
        print(f"[Google Login Error] {e}")
        return jsonify({"success": False, "message": "Invalid Google credentials"}), 401

@app.route("/api/check", methods=["GET"])
def check_login():
    print(f"[DEBUG] Check login called. Cookies: {dict(request.cookies)}")
    username = get_user_from_token()
    if not username:
        print("[DEBUG] No valid user found, returning loggedIn: False")
        return jsonify({"loggedIn": False})
    
    # Check if user logged in via Google (has no password)
    is_google_user = username in users and users[username] == ""
    
    print(f"[DEBUG] Valid user found: {username}, returning loggedIn: True")
    return jsonify({
        "loggedIn": True, 
        "user": username,
        "isGoogleUser": is_google_user
    })

@app.route("/api/logout", methods=["POST"])
def logout():
    resp = jsonify({"success": True, "message": "Logged out"})
    # Clear the JWT cookie by setting it to expire immediately
    secure_flag = False if app.debug else True
    resp.set_cookie("token", "", httponly=True, secure=secure_flag, samesite='Lax', expires=0)
    return resp

# Protected routes (require valid JWT cookie)

@app.route("/api/upload", methods=["POST"])
def upload_file():
    user = get_user_from_token()
    if not user:
        return jsonify({"error": "Unauthorized"}), 401
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400
    filename = file.filename
    
    # Save file using storage manager
    storage_provider = storage_manager.get_user_storage_provider(user)
    
    # Check for duplicate file
    if storage_provider.file_exists(user, filename):
        # Duplicate file upload attempt
        msg_user = {"role": "user", "content": f"📎 Skipped duplicate upload `{filename}`"}
        msg_assistant = {"role": "assistant", "content": "This file is already indexed."}
        # Note: File uploads are global per user, not per session
        conv = get_conversation(user, "default")
        conv.append(msg_user)
        conv.append(msg_assistant)
        save_conversation(user, "default")
        return jsonify({"messages": [msg_user, msg_assistant]})
    
    # Save file to storage
    save_result = storage_provider.save_file(user, filename, file)
    if not save_result:
        # Check if this is a Google Drive authentication issue
        if hasattr(storage_provider, 'real_provider') and storage_provider.real_provider:
            if not storage_provider.real_provider.is_authenticated():
                return jsonify({
                    "error": "Google Drive authentication required",
                    "auth_required": True,
                    "message": "Please complete Google Drive authentication before uploading files"
                }), 401
        
        return jsonify({"error": "Failed to save file"}), 500
    
    # Return immediately with processing message
    msg_user = {"role": "user", "content": f"📎 Uploading `{filename}` for processing..."}
    msg_assistant = {"role": "assistant", "content": f"Processing your file '{filename}'... This may take a few moments for large files."}
    
    # Note: File uploads are global per user, not per session
    conv = get_conversation(user, "default")
    conv.append(msg_user)
    conv.append(msg_assistant)
    save_conversation(user, "default")
    
    # Start background processing
    def process_file_background():
        try:
            print(f"[Background Processing] Starting processing for {filename}")
            
            # Get file from storage for processing
            file_data = storage_provider.get_file(user, filename)
            if not file_data:
                raise Exception("Failed to retrieve saved file")
            
            # Create temporary file for processing
            import tempfile
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{filename.split('.')[-1]}") as temp_file:
                temp_file.write(file_data.read())
                temp_file_path = temp_file.name
            
            file_data.close()
            print(f"[Background Processing] File saved to temp: {temp_file_path}")
            
            # Load or initialize user's indexed files list
            user_dir = os.path.join(books_dir, safe_filename(user))
            os.makedirs(user_dir, exist_ok=True)
            indexed_list_path = os.path.join(user_dir, "indexed_files.json")
            indexed_files = []
            if os.path.exists(indexed_list_path):
                try:
                    with open(indexed_list_path, "r") as idxf:
                        data = json.load(idxf)
                        if isinstance(data, list):
                            indexed_files = data
                        elif isinstance(data, dict) and "indexed_files" in data:
                            indexed_files = data["indexed_files"]
                except:
                    indexed_files = []
            
            print(f"[Background Processing] Processing file type: {filename.split('.')[-1]}")
            
            splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            new_docs = []
            text_blob = ""
            
            ext = filename.lower().rsplit('.', 1)[-1]
            if ext == "pdf":
                print(f"[Background Processing] Processing PDF file")
                loader = PyPDFLoader(temp_file_path)
                pages = loader.load()
                print(f"[Background Processing] PDF loaded with {len(pages)} pages")
                new_docs.extend(splitter.split_documents(pages))
                first_pages_text = [p.page_content for p in pages[:3]]
                text_blob = "\n".join(first_pages_text)
            elif ext in ["jpg", "jpeg", "png"]:
                print(f"[Background Processing] Processing image file with OCR")
                text = pytesseract.image_to_string(Image.open(temp_file_path))
                new_docs.extend(splitter.create_documents([text]))
                text_blob = text
            elif ext in ["mp3", "wav", "m4a"]:
                print(f"[Background Processing] Processing audio file with Whisper")
                model = WhisperModel("base")
                segments, info = model.transcribe(temp_file_path)
                result_text = " ".join([seg.text for seg in segments])
                new_docs.extend(splitter.create_documents([result_text]))
                text_blob = result_text[:2000]
            elif ext == "docx":
                print(f"[Background Processing] Processing Word document")
                doc_obj = docx.Document(temp_file_path)
                full_text = "\n".join([para.text for para in doc_obj.paragraphs])
                new_docs.extend(splitter.create_documents([full_text]))
                text_blob = full_text[:2000]
            elif ext in ["xlsx", "xls"]:
                print(f"[Background Processing] Processing Excel file")
                df = pd.read_excel(temp_file_path, engine="openpyxl" if ext == "xlsx" else "xlrd")
                csv_text = df.to_csv(index=False)
                new_docs.extend(splitter.create_documents([csv_text]))
                text_blob = csv_text[:2000]
            else:
                # Unsupported file type
                raise Exception(f"Unsupported file type: {ext}")
            
            print(f"[Background Processing] Created {len(new_docs)} document chunks")
            
            # If new docs were successfully created, add to vectorstore
            if new_docs:
                print(f"[Background Processing] Adding documents to vectorstore")
                vs = load_vectorstore_for_user(user)
                if vs is None:
                    # Create new FAISS vectorstore for user
                    print(f"[Background Processing] Creating new vectorstore for user")
                    from langchain.vectorstores import FAISS
                    vs = FAISS.from_documents(new_docs, embeddings)
                else:
                    # Add documents to existing vectorstore
                    print(f"[Background Processing] Adding to existing vectorstore")
                    vs.add_documents(new_docs)
                
                if vs:
                    print(f"[Background Processing] Saving vectorstore")
                    save_vectorstore_for_user(user, vs)
                    print(f"[Background Processing] Vectorstore saved successfully")
            
            # Update indexed files list on disk (only for local storage)
            if not isinstance(storage_provider, GoogleDriveStorageProvider):
                indexed_files.append(filename)
                try:
                    with open(indexed_list_path, 'w') as idxf:
                        json.dump({"indexed_files": indexed_files}, idxf)
                except Exception as e:
                    print(f"[Indexed file list save error] {e}")
            else:
                # For Google Drive (Demo), files are automatically "indexed" when stored
                print(f"[GoogleDrive] File {filename} automatically indexed in demo storage")
            
            # Clean up temporary file
            try:
                os.unlink(temp_file_path)
            except Exception as e:
                print(f"[Temp file cleanup error] {e}")
            
            # Update chat with success message
            print(f"[Background Processing] Processing completed successfully")
            summary = ""
            if text_blob:
                snippet = (text_blob[:500] + "...") if len(text_blob) > 500 else text_blob
                summary = f" Here's a snippet of the content: \n{snippet}"
            
            success_msg = {"role": "assistant", "content": f"✅ Your file '{filename}' has been successfully indexed and is ready to use in your studies!{summary}"}
            
            conv = get_conversation(user, "default")
            # Replace the processing message with success message
            if conv and conv[-1]["role"] == "assistant" and "Processing your file" in conv[-1]["content"]:
                conv[-1] = success_msg
            else:
                conv.append(success_msg)
            save_conversation(user, "default")
            print(f"[Background Processing] Success message updated in chat")
            
        except Exception as e:
            print(f"[Background Processing] Error processing file: {e}")
            print(f"[Background Processing] Full error details: {type(e).__name__}: {str(e)}")
            # Update chat with error message
            error_msg = {"role": "assistant", "content": f"❌ Sorry, I couldn't process the file '{filename}'. Please try again or contact support."}
            
            conv = get_conversation(user, "default")
            # Replace the processing message with error message
            if conv and conv[-1]["role"] == "assistant" and "Processing your file" in conv[-1]["content"]:
                conv[-1] = error_msg
            else:
                conv.append(error_msg)
            save_conversation(user, "default")
            print(f"[Background Processing] Error message updated in chat")
    
    # Start background processing thread
    thread = threading.Thread(target=process_file_background)
    thread.daemon = True
    thread.start()
    
    return jsonify({"messages": [msg_user, msg_assistant]})

@app.route("/api/upload/status", methods=["GET"])
def get_upload_status():
    """Get upload processing status for the current user"""
    user = get_user_from_token()
    if not user:
        return jsonify({"error": "Unauthorized"}), 401
    
    # Get the latest conversation to check processing status
    conv = get_conversation(user, "default")
    if not conv:
        return jsonify({"status": "no_uploads"})
    
    # Check the last assistant message
    last_assistant_msg = None
    for msg in reversed(conv):
        if msg["role"] == "assistant":
            last_assistant_msg = msg
            break
    
    if not last_assistant_msg:
        return jsonify({"status": "no_uploads"})
    
    content = last_assistant_msg["content"]
    
    if "Processing your file" in content:
        return jsonify({"status": "processing"})
    elif "✅ Your file" in content and "has been successfully indexed" in content:
        return jsonify({"status": "completed"})
    elif "❌ Sorry, I couldn't process" in content:
        return jsonify({"status": "failed"})
    else:
        return jsonify({"status": "unknown"})

@app.route("/api/chat", methods=["POST"])
def chat():
    user = get_user_from_token()
    if not user:
        return jsonify({"error": "Unauthorized"}), 401
    data = request.get_json()
    message = (data.get("message") or "").strip()
    session_id = data.get("session_id", "default")
    if not message:
        return jsonify({"error": "No message provided"}), 400
    conv = get_conversation(user, session_id)
    profile = extract_user_profile(conv, default_name=user.split('@')[0].capitalize())
    user_name = profile["name"]
    interests = profile["interests"]
    # Greeting handling
    if message.lower() in ["hi", "hello", "hey", "hi there", "hello there"]:
        greeting = f"Hi {user_name}! 👋 How can I assist you today?"
        assistant_msg = {"role": "assistant", "content": greeting}
        conv.append({"role": "user", "content": message})
        conv.append(assistant_msg)
        save_conversation(user)
        return jsonify({"messages": [assistant_msg]})
    # Construct context prefix from user profile
    context_prefix = f"User Name: {user_name}\n"
    if interests:
        context_prefix += "User Interests: " + ", ".join(interests) + "\n"
    # Recent chat history snippet (last 6 messages) for context
    history_snippets = []
    for msg in conv[-6:]:
        if msg["role"] == "user":
            history_snippets.append(f"User: {msg['content']}")
        elif msg["role"] == "assistant":
            history_snippets.append(f"Assistant: {msg['content']}")
    chat_history_str = "\n".join(history_snippets)
    # Retrieve relevant docs for the query (user's documents + global knowledge)
    vs = load_vectorstore_for_user(user)
    docs = []
    
    # Search user's personal documents
    if vs is not None:
        query = f"{context_prefix}Question: {message}"
        try:
            user_docs = vs.similarity_search(query, k=3)
            docs.extend(user_docs)
        except Exception as e:
            print(f"[User VectorStore Error] {e}")
    
    # Search global knowledge base
    try:
        global_docs = search_global_knowledge(message, k=2)
        docs.extend(global_docs)
    except Exception as e:
        print(f"[Global Knowledge Search Error] {e}")
    
    # Remove duplicates and limit total docs
    unique_docs = []
    seen_content = set()
    for doc in docs:
        if doc.page_content not in seen_content:
            unique_docs.append(doc)
            seen_content.add(doc.page_content)
        if len(unique_docs) >= 5:  # Limit to 5 total docs
            break
    docs = unique_docs
    # Build LLM prompt with context if available
    if docs:
        docs_text = "\n".join([doc.page_content for doc in docs])
        system_content = (
            "You are a knowledgeable teacher assistant. You strictly rely on the provided content to answer the question.\n"
            "If the context does NOT contain enough information, politely say you couldn't find relevant info in the material, and then give a brief general explanation.\n"
            "You can understand and respond in English, Tamil, or Arabic as appropriate.\n"
            f"Context:\n{docs_text}\n\n"
            f"Chat History:\n{chat_history_str}"
        )
    else:
        system_content = (
            "You are a helpful teacher assistant. Answer the user's question clearly and truthfully.\n"
            "If the question refers to uploaded documents but no relevant info is found, apologize for not finding info in the material and answer generally.\n"
            "You can understand and respond in English, Tamil, or Arabic as appropriate."
        )
    user_content = f"{context_prefix}Question: {message}"
    messages = [
        {"role": "system", "content": system_content},
        {"role": "user", "content": user_content}
    ]
    # Check if API key is available
    if not THETA_API_KEY:
        answer_text = "⚠️ LLM API key not configured. Please set THETA_API_KEY in your environment."
    else:
        result = llm.invoke(messages)
        if isinstance(result, str):
            answer_text = result
        elif isinstance(result, dict):
            answer_text = result.get("content", "")
        else:
            answer_text = str(result)
    
    assistant_msg = {"role": "assistant", "content": answer_text}
    conv.append({"role": "user", "content": message})
    conv.append(assistant_msg)
    save_conversation(user, session_id)
    
    # Update session timestamp
    sessions = get_user_sessions(user)
    for session in sessions:
        if session["id"] == session_id:
            session["updated_at"] = datetime.datetime.now().isoformat()
            break
    save_user_sessions(user)
    
    return jsonify({"messages": [assistant_msg]})



@app.route("/api/history", methods=["GET"])
def get_history():
    user = get_user_from_token()
    if not user:
        return jsonify({"error": "Unauthorized"}), 401
    session_id = request.args.get("session_id", "default")
    conv = get_conversation(user, session_id)
    return jsonify({"messages": conv[-10:]})

@app.route("/api/sessions", methods=["GET"])
def get_sessions():
    """Get all chat sessions for the current user"""
    user = get_user_from_token()
    if not user:
        return jsonify({"error": "Unauthorized"}), 401
    sessions = get_user_sessions(user)
    return jsonify({"sessions": sessions})

@app.route("/api/sessions", methods=["POST"])
def create_session():
    """Create a new chat session"""
    user = get_user_from_token()
    if not user:
        return jsonify({"error": "Unauthorized"}), 401
    data = request.get_json()
    name = data.get("name", "New Chat")
    new_session = create_new_session(user, name)
    return jsonify({"session": new_session})

@app.route("/api/sessions/<session_id>/rename", methods=["PUT"])
def rename_session(session_id):
    """Rename a chat session"""
    user = get_user_from_token()
    if not user:
        return jsonify({"error": "Unauthorized"}), 401
    data = request.get_json()
    new_name = data.get("name", "New Chat")
    success = update_session_name(user, session_id, new_name)
    if success:
        return jsonify({"success": True, "message": "Session renamed"})
    else:
        return jsonify({"error": "Session not found"}), 404

@app.route("/api/sessions/<session_id>", methods=["DELETE"])
def delete_session_endpoint(session_id):
    """Delete a chat session"""
    user = get_user_from_token()
    if not user:
        return jsonify({"error": "Unauthorized"}), 401
    success = delete_session(user, session_id)
    if success:
        return jsonify({"success": True, "message": "Session deleted"})
    else:
        return jsonify({"error": "Session not found"}), 404

@app.route("/api/test-llm", methods=["POST"])
def test_llm():
    """Test endpoint to verify LLM functionality"""
    try:
        test_messages = [
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": "Say hello in one sentence."}
        ]
        result = llm.invoke(test_messages)
        return jsonify({"success": True, "response": result})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/global-knowledge", methods=["POST"])
def update_global_knowledge():
    """Add knowledge to the global vectorstore"""
    user = get_user_from_token()
    if not user:
        return jsonify({"error": "Unauthorized"}), 401
    
    data = request.get_json()
    texts = data.get("texts", [])
    
    if not texts:
        return jsonify({"error": "No texts provided"}), 400
    
    try:
        update_global_vectorstore(texts)
        return jsonify({"success": True, "message": f"Added {len(texts)} texts to global knowledge base"})
    except Exception as e:
        print(f"[Global Knowledge Update Error] {e}")
        return jsonify({"error": "Failed to update global knowledge"}), 500

@app.route("/api/global-knowledge", methods=["GET"])
def get_global_knowledge_info():
    """Get information about the global knowledge base"""
    try:
        global global_vectorstore
        if not global_vectorstore:
            global_vectorstore = load_global_vectorstore()
        
        # Get basic info about the vectorstore
        if hasattr(global_vectorstore, 'index_to_docstore_id'):
            doc_count = len(global_vectorstore.index_to_docstore_id)
        else:
            doc_count = 0
        
        return jsonify({
            "success": True,
            "documents_count": doc_count,
            "status": "active"
        })
    except Exception as e:
        print(f"[Global Knowledge Info Error] {e}")
        return jsonify({"error": "Failed to get global knowledge info"}), 500

@app.route("/api/documents", methods=["GET"])
def get_documents():
    """Get list of user's uploaded documents"""
    user = get_user_from_token()
    if not user:
        return jsonify({"error": "Unauthorized"}), 401
    
    try:
        documents = get_user_documents(user)
        return jsonify({
            "success": True,
            "documents": documents,
            "total_count": len(documents)
        })
    except Exception as e:
        print(f"[Documents List Error] {e}")
        return jsonify({"error": "Failed to get documents list"}), 500

@app.route("/api/documents/<filename>", methods=["DELETE"])
def delete_document(filename):
    """Delete a specific document"""
    user = get_user_from_token()
    if not user:
        return jsonify({"error": "Unauthorized"}), 401
    
    try:
        success, message = delete_user_document(user, filename)
        if success:
            return jsonify({
                "success": True,
                "message": message
            })
        else:
            return jsonify({"error": message}), 400
    except Exception as e:
        print(f"[Document Deletion Error] {e}")
        return jsonify({"error": "Failed to delete document"}), 500

@app.route("/api/documents/clear-vectorstore", methods=["DELETE"])
def clear_user_vectorstore():
    """Clear user's entire vectorstore (delete all indexed documents)"""
    user = get_user_from_token()
    if not user:
        return jsonify({"error": "Unauthorized"}), 401
    
    try:
        print(f"[Vectorstore Clear] User {user} requested vectorstore clear")
        
        # Clear vectorstore from cache
        if user in vectorstores_cache:
            del vectorstores_cache[user]
            print(f"[Vectorstore Clear] Cleared from cache for user {user}")
        
        # Delete vectorstore files
        user_vector_dir = os.path.join(vectorstores_dir, safe_filename(user))
        if os.path.exists(user_vector_dir):
            import shutil
            shutil.rmtree(user_vector_dir)
            print(f"[Vectorstore Clear] Deleted vectorstore directory for user {user}")
        
        # Clear indexed files list (for local storage compatibility)
        user_dir = os.path.join(books_dir, safe_filename(user))
        indexed_list_path = os.path.join(user_dir, "indexed_files.json")
        if os.path.exists(indexed_list_path):
            try:
                with open(indexed_list_path, 'w') as idxf:
                    json.dump({"indexed_files": []}, idxf)
                print(f"[Vectorstore Clear] Cleared indexed files list for user {user}")
            except Exception as e:
                print(f"[Vectorstore Clear] Error clearing indexed files list: {e}")
        
        # Add confirmation message to chat
        clear_msg = {"role": "assistant", "content": "🗑️ Your entire document knowledge base has been cleared. All indexed documents have been removed from your vectorstore. You can upload new documents to rebuild your knowledge base."}
        
        conv = get_conversation(user, "default")
        conv.append(clear_msg)
        save_conversation(user, "default")
        
        print(f"[Vectorstore Clear] Successfully cleared vectorstore for user {user}")
        return jsonify({
            "success": True,
            "message": "Vectorstore cleared successfully. All indexed documents have been removed."
        })
        
    except Exception as e:
        print(f"[Vectorstore Clear Error] {e}")
        return jsonify({"error": "Failed to clear vectorstore"}), 500

@app.route("/api/documents/search", methods=["POST"])
def search_documents():
    """Search user's documents with a query"""
    user = get_user_from_token()
    if not user:
        return jsonify({"error": "Unauthorized"}), 401
    
    data = request.get_json()
    query = data.get("query", "").strip()
    k = data.get("k", 5)
    
    if not query:
        return jsonify({"error": "No query provided"}), 400
    
    try:
        results = search_user_documents(user, query, k)
        
        # Format results for frontend
        formatted_results = []
        for doc in results:
            formatted_results.append({
                "content": doc.page_content,
                "metadata": doc.metadata if hasattr(doc, 'metadata') else {},
                "score": getattr(doc, 'score', None)
            })
        
        return jsonify({
            "success": True,
            "query": query,
            "results": formatted_results,
            "total_found": len(formatted_results)
        })
    except Exception as e:
        print(f"[Document Search Error] {e}")
        return jsonify({"error": "Failed to search documents"}), 500

@app.route("/api/documents/stats", methods=["GET"])
def get_document_stats():
    """Get document statistics for user"""
    user = get_user_from_token()
    if not user:
        return jsonify({"error": "Unauthorized"}), 401
    
    try:
        documents = get_user_documents(user)
        
        # Calculate statistics
        total_size = sum(doc["size"] for doc in documents)
        file_types = {}
        for doc in documents:
            file_type = doc["file_type"]
            file_types[file_type] = file_types.get(file_type, 0) + 1
        
        stats = {
            "total_documents": len(documents),
            "total_size_bytes": total_size,
            "total_size_mb": round(total_size / (1024 * 1024), 2),
            "file_types": file_types,
            "supported_types": ["pdf", "docx", "xlsx", "xls", "jpg", "jpeg", "png", "mp3", "wav", "m4a"]
        }
        
        return jsonify({
            "success": True,
            "stats": stats
        })
    except Exception as e:
        print(f"[Document Stats Error] {e}")
        return jsonify({"error": "Failed to get document statistics"}), 500

@app.route("/api/storage/preferences", methods=["GET"])
def get_storage_preferences():
    """Get user's storage preferences"""
    user = get_user_from_token()
    if not user:
        return jsonify({"error": "Unauthorized"}), 401
    
    try:
        current_provider = storage_manager.user_preferences.get(user, 'local')
        # Don't call get_available_providers() here to avoid triggering Google Drive init
        available_providers = ['local', 'google_drive']  # Hardcode for now
        
        # Check if user logged in via Google
        is_google_user = user in users and users[user] == ""
        
        return jsonify({
            "success": True,
            "current_provider": current_provider,
            "available_providers": available_providers,
            "isGoogleUser": is_google_user
        })
    except Exception as e:
        print(f"[Storage Preferences Error] {e}")
        return jsonify({"error": "Failed to get storage preferences"}), 500

@app.route("/api/storage/preferences", methods=["POST"])
def set_storage_preferences():
    """Set user's storage preferences"""
    user = get_user_from_token()
    if not user:
        return jsonify({"error": "Unauthorized"}), 401
    
    data = request.get_json()
    provider = data.get("provider")
    
    if not provider:
        return jsonify({"error": "No provider specified"}), 400
    
    try:
        storage_manager.set_user_storage_preference(user, provider)
        return jsonify({
            "success": True,
            "message": f"Storage preference set to {provider}"
        })
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        print(f"[Storage Preferences Error] {e}")
        return jsonify({"error": "Failed to set storage preferences"}), 500

@app.route("/api/storage/migrate", methods=["POST"])
def migrate_storage():
    """Migrate user data between storage providers"""
    user = get_user_from_token()
    if not user:
        return jsonify({"error": "Unauthorized"}), 401
    
    data = request.get_json()
    from_provider = data.get("from_provider")
    to_provider = data.get("to_provider")
    
    if not from_provider or not to_provider:
        return jsonify({"error": "Both from_provider and to_provider required"}), 400
    
    try:
        success = storage_manager.migrate_user_data(user, from_provider, to_provider)
        if success:
            return jsonify({
                "success": True,
                "message": f"Successfully migrated from {from_provider} to {to_provider}"
            })
        else:
            return jsonify({"error": "Migration failed"}), 500
    except Exception as e:
        print(f"[Storage Migration Error] {e}")
        return jsonify({"error": "Failed to migrate data"}), 500

@app.route("/api/storage/google-drive/auth", methods=["GET"])
def get_google_drive_auth_url():
    """Get Google Drive authorization URL"""
    global used_auth_codes_global
    
    try:
        # Get current user from token
        user = get_user_from_token()
        if not user:
            return jsonify({"error": "Authentication required"}), 401
        
        # Clear global used auth codes when getting a new auth URL
        print(f"[OAuth] 🧹 Clearing global used auth codes for fresh authentication")
        used_auth_codes_global.clear()
        
        # Check if credentials file exists
        credentials_exist = os.path.exists("credentials.json")
        
        if not credentials_exist:
            return jsonify({
                "success": False,
                "error": "Google Drive credentials not found. Please set up credentials.json first."
            }), 400
        
        # Initialize Google Drive provider only when needed
        google_drive_provider = storage_manager.providers.get('google_drive')
        if google_drive_provider is None:
            storage_manager.providers['google_drive'] = GoogleDriveStorageProvider()
            google_drive_provider = storage_manager.providers['google_drive']
        
        if google_drive_provider.is_authenticated(user):
            return jsonify({
                "success": True,
                "authenticated": True,
                "message": "Google Drive is already authenticated for this user"
            })
        
        auth_url = google_drive_provider.get_auth_url(user)
        if auth_url:
            return jsonify({
                "success": True,
                "authenticated": False,
                "auth_url": auth_url,
                "message": "Please visit the authorization URL to connect Google Drive"
            })
        else:
            return jsonify({
                "success": False,
                "error": "Google Drive service not available. Please try again later."
            }), 400
            
    except Exception as e:
        print(f"[Google Drive Auth Error] {e}")
        return jsonify({"error": "Failed to get authorization URL"}), 500

@app.route("/api/storage/google-drive/auth", methods=["POST"])
def complete_google_drive_auth():
    """Complete Google Drive OAuth flow"""
    try:
        # Get current user from token
        user = get_user_from_token()
        if not user:
            return jsonify({"error": "Authentication required"}), 401
        
        data = request.get_json()
        auth_code = data.get("auth_code")
        
        if not auth_code:
            return jsonify({"error": "Authorization code required"}), 400
        
        google_drive_provider = storage_manager.providers.get('google_drive')
        if not google_drive_provider:
            return jsonify({"error": "Google Drive provider not available"}), 500
        
        success = google_drive_provider.complete_auth(auth_code, user)
        if success:
            return jsonify({
                "success": True,
                "message": "Google Drive authentication completed successfully for this user"
            })
        else:
            return jsonify({"error": "Authentication failed"}), 500
            
    except Exception as e:
        print(f"[Google Drive Auth Error] {e}")
        return jsonify({"error": "Failed to complete authentication"}), 500

# Global set to track used authorization codes across all requests
used_auth_codes_global = set()

@app.route("/oauth2callback")
def oauth2callback():
    """Handle Google OAuth callback"""
    global used_auth_codes_global
    
    try:
        print(f"[OAuth] 🔄 OAuth callback received")
        print(f"[OAuth] 📋 Request args: {dict(request.args)}")
        
        # Get the authorization code from the request
        auth_code = request.args.get('code')
        if not auth_code:
            print(f"[OAuth] ❌ No authorization code received")
            return redirect("http://localhost:3000?auth=error&message=No authorization code received")
        
        print(f"[OAuth] 🔑 Authorization code received: {auth_code[:20]}...")
        
        # Check if this auth code has already been used globally
        if auth_code in used_auth_codes_global:
            print(f"[OAuth] 🚫 Authorization code already used globally: {auth_code[:20]}...")
            return redirect("http://localhost:3000?auth=error&message=Authorization code already used")
        
        # Mark this auth code as used immediately
        used_auth_codes_global.add(auth_code)
        print(f"[OAuth] 🔒 Marked auth code as used: {auth_code[:20]}...")
        
        # Get the state parameter which should contain the user info
        state = request.args.get('state')
        user = None
        
        print(f"[OAuth] 📝 State parameter received: {state}")
        
        if state:
            try:
                # Decode user from state parameter
                import base64
                user = base64.b64decode(state.encode()).decode()
                print(f"[OAuth] 👤 User extracted from state: {user}")
            except Exception as e:
                print(f"[OAuth] ❌ Error decoding state: {e}")
                print(f"[OAuth] 🔍 State parameter that failed to decode: {state}")
                pass
        
        # If no user in state, we can't complete authentication
        if not user:
            print(f"[OAuth] ❌ No user information found in state")
            return redirect("http://localhost:3000?auth=error&message=No user information found")
        
        print(f"[OAuth] 🔐 Starting authentication completion for user: {user}")
        
        # Complete the authentication for this specific user
        google_drive_provider = storage_manager.providers.get('google_drive')
        if not google_drive_provider:
            print(f"[OAuth] 🔧 Creating new Google Drive provider")
            storage_manager.providers['google_drive'] = GoogleDriveStorageProvider()
            google_drive_provider = storage_manager.providers['google_drive']
        
        print(f"[OAuth] 🔄 Calling complete_auth for user: {user}")
        success = google_drive_provider.complete_auth(auth_code, user)
        
        if success:
            print(f"[OAuth] ✅ Authentication successful for user: {user}")
            # Redirect back to the frontend with success message
            return redirect(f"http://localhost:3000?auth=success&user={user}")
        else:
            print(f"[OAuth] ❌ Authentication failed for user: {user}")
            # Redirect back to the frontend with error message
            return redirect(f"http://localhost:3000?auth=error&user={user}")
            
    except Exception as e:
        print(f"[OAuth] 💥 OAuth callback error: {e}")
        print(f"[OAuth] 🔍 Full error details: {type(e).__name__}: {str(e)}")
        return redirect("http://localhost:3000?auth=error")

@app.route("/oauth2callback/")
def oauth2callback_alt():
    """Alternative OAuth callback endpoint"""
    return oauth2callback()

@app.route("/api/storage/google-drive/callback")
def google_drive_callback():
    """Handle Google Drive OAuth callback"""
    return oauth2callback()

@app.route("/api/storage/google-drive/status", methods=["GET"])
def get_google_drive_status():
    """Get Google Drive authentication status for current user"""
    try:
        # Get current user from token
        user = get_user_from_token()
        if not user:
            return jsonify({"error": "Authentication required"}), 401
        
        # Check if credentials file exists
        credentials_exist = os.path.exists("credentials.json")
        
        if not credentials_exist:
            return jsonify({
                "success": True,
                "authenticated": False,
                "credentials_available": False,
                "message": "Google Drive credentials not found"
            })
        
        # Initialize Google Drive provider to check actual authentication status
        google_drive_provider = storage_manager.providers.get('google_drive')
        if google_drive_provider is None:
            storage_manager.providers['google_drive'] = GoogleDriveStorageProvider()
            google_drive_provider = storage_manager.providers['google_drive']
        
        is_authenticated = google_drive_provider.is_authenticated(user)
        
        # Get auth URL if not authenticated
        auth_url = None
        if not is_authenticated:
            auth_url = google_drive_provider.get_auth_url(user)
        
        return jsonify({
            "success": True,
            "authenticated": is_authenticated,
            "credentials_available": credentials_exist,
            "message": "Google Drive is authenticated for this user" if is_authenticated else "Google Drive needs authentication for this user",
            "auth_url": auth_url
        })
        
    except Exception as e:
        print(f"[Google Drive Status Error] {e}")
        return jsonify({"error": "Failed to get status"}), 500

@app.route("/api/audio", methods=["POST"])
def audio_question():
    user = get_user_from_token()
    if not user:
        return jsonify({"error": "Unauthorized"}), 401
    if "file" not in request.files:
        return jsonify({"error": "No audio file provided"}), 400
    audio_file = request.files["file"]
    if audio_file.filename == "":
        return jsonify({"error": "No audio file provided"}), 400
    # Save audio to a temp path
    temp_audio_path = os.path.join(books_dir, safe_filename(user), "temp_audio_input")
    os.makedirs(os.path.dirname(temp_audio_path), exist_ok=True)
    audio_file.save(temp_audio_path)
    try:
        model = WhisperModel("base")
        segments, info = model.transcribe(temp_audio_path)
        question_text = " ".join([seg.text for seg in segments]).strip()
    except Exception as e:
        print(f"[Whisper Error] {e}")
        return jsonify({"error": "Audio transcription failed"}), 500
    finally:
        try:
            os.remove(temp_audio_path)
        except:
            pass
    if not question_text:
        return jsonify({"error": "Unable to transcribe audio"}), 400
    conv = get_conversation(user)
    profile = extract_user_profile(conv, default_name=user.split('@')[0].capitalize())
    user_name = profile["name"]
    interests = profile["interests"]
    # Greeting via audio
    if question_text.lower() in ["hi", "hello", "hey", "hi there", "hello there"]:
        greeting = f"Hi {user_name}! 👋 How can I assist you today?"
        assistant_msg = {"role": "assistant", "content": greeting}
        conv.append({"role": "user", "content": question_text})
        conv.append(assistant_msg)
        save_conversation(user, "default")  # Audio uses default session
        return jsonify({"messages": [{"role": "user", "content": question_text}, assistant_msg]})
    context_prefix = f"User Name: {user_name}\n"
    if interests:
        context_prefix += "User Interests: " + ", ".join(interests) + "\n"
    history_snippets = []
    for msg in conv[-6:]:
        if msg["role"] == "user":
            history_snippets.append(f"User: {msg['content']}")
        elif msg["role"] == "assistant":
            history_snippets.append(f"Assistant: {msg['content']}")
    chat_history_str = "\n".join(history_snippets)
    # Retrieve relevant docs for the query (user's documents + global knowledge)
    vs = load_vectorstore_for_user(user)
    docs = []
    
    # Search user's personal documents
    if vs is not None:
        query = f"{context_prefix}Question: {question_text}"
        try:
            user_docs = vs.similarity_search(query, k=3)
            docs.extend(user_docs)
        except Exception as e:
            print(f"[User VectorStore Error] {e}")
    
    # Search global knowledge base
    try:
        global_docs = search_global_knowledge(question_text, k=2)
        docs.extend(global_docs)
    except Exception as e:
        print(f"[Global Knowledge Search Error] {e}")
    
    # Remove duplicates and limit total docs
    unique_docs = []
    seen_content = set()
    for doc in docs:
        if doc.page_content not in seen_content:
            unique_docs.append(doc)
            seen_content.add(doc.page_content)
        if len(unique_docs) >= 5:  # Limit to 5 total docs
            break
    docs = unique_docs
    if docs:
        docs_text = "\n".join([doc.page_content for doc in docs])
        system_content = (
            "You are a knowledgeable teacher assistant. You strictly rely on the provided content to answer the question.\n"
            "If the context does NOT contain enough information, politely say you couldn't find relevant info in the material, and then provide a brief general explanation.\n"
            "You can understand and respond in English, Tamil, or Arabic as appropriate.\n"
            f"Context:\n{docs_text}\n\n"
            f"Chat History:\n{chat_history_str}"
        )
    else:
        system_content = (
            "You are a helpful teacher assistant. Answer the user's question clearly and truthfully.\n"
            "If the question refers to the uploaded documents but no relevant info is found, apologize for not finding info and answer generally.\n"
            "You can understand and respond in English, Tamil, or Arabic as appropriate."
        )
    user_content = f"User Name: {user_name}\nQuestion: {question_text}"
    messages = [
        {"role": "system", "content": system_content},
        {"role": "user", "content": user_content}
    ]
    # Check if API key is available
    if not THETA_API_KEY:
        answer_text = "⚠️ LLM API key not configured. Please set THETA_API_KEY in your environment."
    else:
        result = llm.invoke(messages)
        if isinstance(result, str):
            answer_text = result
        elif isinstance(result, dict):
            answer_text = result.get("content", "")
        else:
            answer_text = str(result)
    
    assistant_msg = {"role": "assistant", "content": answer_text}
    conv.append({"role": "user", "content": question_text})
    conv.append(assistant_msg)
    save_conversation(user, "default")  # Audio uses default session
    return jsonify({"messages": [{"role": "user", "content": question_text}, assistant_msg]})

# Initialize global vectorstore on startup
print("[STARTUP] Initializing global vectorstore...")
try:
    load_global_vectorstore()
    print("[STARTUP] Global vectorstore initialized successfully!")
except Exception as e:
    print(f"[STARTUP] Global vectorstore initialization failed: {e}")
    print("[STARTUP] Continuing without global vectorstore...")

if __name__ == "__main__":
    # Production-ready configuration
    port = int(os.getenv("PORT", 5000))
    host = os.getenv("HOST", "0.0.0.0")
    debug = os.getenv("FLASK_DEBUG", "False").lower() == "true"
    
    print(f"[STARTUP] Starting Flask app on {host}:{port}")
    print(f"[STARTUP] Debug mode: {debug}")
    print(f"[STARTUP] Environment: {os.getenv('FLASK_ENV', 'development')}")
    
    app.run(host=host, port=port, debug=debug)

